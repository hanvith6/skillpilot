from fastapi import FastAPI, APIRouter, HTTPException, Depends, Request, Header
from fastapi.responses import FileResponse, StreamingResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict, EmailStr
from typing import List, Optional, Dict, Any
import uuid
from datetime import datetime, timezone, timedelta
import jwt
from passlib.context import CryptContext
from google import genai
import razorpay
from emergentintegrations.payments.stripe.checkout import StripeCheckout, CheckoutSessionRequest, CheckoutSessionResponse, CheckoutStatusResponse
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# Initialize services
client = genai.Client(api_key=os.environ['GEMINI_API_KEY'])
razorpay_client = razorpay.Client(auth=(os.environ['RAZORPAY_KEY_ID'], os.environ['RAZORPAY_KEY_SECRET']))

# Password hashing
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

# JWT settings
JWT_SECRET = os.environ['JWT_SECRET']
JWT_ALGORITHM = "HS256"
JWT_EXPIRATION = timedelta(days=30)

# Create the main app
app = FastAPI()
api_router = APIRouter(prefix="/api")

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Models
class UserSignup(BaseModel):
    name: str
    email: EmailStr
    password: str
    referral_code: Optional[str] = None

class UserLogin(BaseModel):
    email: EmailStr
    password: str

class User(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str
    email: EmailStr
    credits: int = 100
    referral_code: str = Field(default_factory=lambda: f"SKILLMATE-{str(uuid.uuid4())[:8].upper()}")
    referred_by: Optional[str] = None
    total_referrals: int = 0
    referral_credits_earned: int = 0
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class ReferralStats(BaseModel):
    referral_code: str
    total_referrals: int
    credits_earned: int
    recent_referrals: List[Dict[str, Any]]

class AuthResponse(BaseModel):
    token: str
    user: User

class ResumeRequest(BaseModel):
    resume_text: str
    target_role: str
    country: str
    emergent_mode: bool = False

class ProjectRequest(BaseModel):
    topic: str
    branch: str
    emergent_mode: bool = False

class EnglishRequest(BaseModel):
    text: str
    emergent_mode: bool = False

class InterviewRequest(BaseModel):
    question_type: str
    background: Optional[str] = None
    emergent_mode: bool = False

class GenerationHistory(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    user_id: str
    type: str
    title: str
    content: Dict[str, Any]
    credits_used: int
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class PurchaseRequest(BaseModel):
    package_id: str

class RazorpayOrderResponse(BaseModel):
    order_id: str
    amount: int
    currency: str
    key_id: str

class RazorpayVerifyRequest(BaseModel):
    order_id: str
    payment_id: str
    signature: str

class StripeCheckoutRequest(BaseModel):
    package_id: str

class PaymentTransaction(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    user_id: str
    session_id: Optional[str] = None
    payment_id: Optional[str] = None
    amount: float
    currency: str
    package_id: str
    credits: int
    status: str
    payment_status: str
    provider: str
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

# Credit packages
CREDIT_PACKAGES = {
    "starter_inr": {"credits": 100, "price": 99, "currency": "INR", "region": "india"},
    "pro_inr": {"credits": 300, "price": 249, "currency": "INR", "region": "india"},
    "unlimited_inr": {"credits": 600, "price": 299, "currency": "INR", "region": "india"},
    "starter_usd": {"credits": 100, "price": 4.0, "currency": "USD", "region": "global"},
    "pro_usd": {"credits": 300, "price": 5.0, "currency": "USD", "region": "global"},
    "unlimited_usd": {"credits": 600, "price": 6.0, "currency": "USD", "region": "global"},
}

# Credit costs
CREDIT_COSTS = {
    "resume": 20,
    "project": 25,
    "english": 10,
    "interview": 15,
}

# Auth helpers
def verify_password(plain_password, hashed_password):
    return pwd_context.verify(plain_password, hashed_password)

def get_password_hash(password):
    return pwd_context.hash(password)

def create_access_token(data: dict):
    to_encode = data.copy()
    expire = datetime.now(timezone.utc) + JWT_EXPIRATION
    to_encode.update({"exp": expire})
    return jwt.encode(to_encode, JWT_SECRET, algorithm=JWT_ALGORITHM)

async def get_current_user(request: Request, authorization: str = Header(None)):
    # First check session_token cookie (Google OAuth)
    session_token = request.cookies.get("session_token")
    
    if session_token:
        session_doc = await db.user_sessions.find_one({"session_token": session_token}, {"_id": 0})
        if session_doc:
            expires_at = session_doc["expires_at"]
            if isinstance(expires_at, str):
                expires_at = datetime.fromisoformat(expires_at)
            if expires_at.tzinfo is None:
                expires_at = expires_at.replace(tzinfo=timezone.utc)
            
            if expires_at < datetime.now(timezone.utc):
                raise HTTPException(status_code=401, detail="Session expired")
            
            user_doc = await db.users.find_one({"id": session_doc["user_id"]}, {"_id": 0})
            if user_doc:
                if isinstance(user_doc.get('created_at'), str):
                    user_doc['created_at'] = datetime.fromisoformat(user_doc['created_at'])
                return User(**user_doc)
    
    # Fallback to JWT token (email/password auth)
    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Not authenticated")
    
    token = authorization.split(" ")[1]
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        user_id = payload.get("sub")
        if user_id is None:
            raise HTTPException(status_code=401, detail="Invalid token")
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")
    
    user_doc = await db.users.find_one({"id": user_id}, {"_id": 0})
    if not user_doc:
        raise HTTPException(status_code=404, detail="User not found")
    
    if isinstance(user_doc.get('created_at'), str):
        user_doc['created_at'] = datetime.fromisoformat(user_doc['created_at'])
    
    return User(**user_doc)

# Gemini prompts
RESUME_PROMPT = """You are an expert resume writer specializing in ATS-optimized resumes for engineering students and fresh graduates.

Target Role: {target_role}
Target Country: {country}

Original Resume/Information:
{resume_text}

Task:
1. Create an ATS-optimized professional resume
2. Suggest relevant technical skills for the role
3. Write a compelling professional summary
4. Format in a clean, parseable structure
5. Use action verbs and quantifiable achievements
6. Tailor for {country} job market standards

Provide the output in JSON format with these keys:
{{
  "summary": "professional summary",
  "skills": ["skill1", "skill2", ...],
  "experience": [{{"title": "", "company": "", "duration": "", "points": []}}],
  "education": [{{"degree": "", "institution": "", "year": ""}}],
  "projects": [{{"name": "", "description": "", "tech": []}}]
}}"""

PROJECT_PROMPT = """You are a senior computer science professor helping students with final year project documentation.

Topic: {topic}
Branch: {branch}

Task: Generate a complete final year project report structure including:

1. Abstract (150-200 words)
2. Problem Statement
3. Objectives (3-5 points)
4. System Architecture overview
5. Modules/Components (4-6 modules with descriptions)
6. Implementation highlights
7. Technology Stack
8. Future Scope (3-4 points)
9. 20 Viva Questions with detailed answers

Make it university-appropriate, technically sound, and plagiarism-free. Use proper academic language.

Provide output in JSON format with keys: abstract, problem_statement, objectives, architecture, modules, implementation, tech_stack, future_scope, viva_questions"""

ENGLISH_PROMPT = """You are an English language expert helping non-native speakers improve their professional communication.

Original Text:
{text}

Task: Rewrite the text in THREE versions:
1. FORMAL: Highly professional, suitable for corporate emails and official documents
2. SEMI_FORMAL: Balanced, suitable for team communication and general business use
3. SIMPLE: Clear and direct, suitable for quick messages and casual professional settings

Rules:
- Maintain the original meaning and intent
- Sound natural and human (avoid AI buzzwords like "delve", "leverage", "synergy")
- Fix grammar and spelling errors
- Keep it concise

Provide output in JSON format with keys: formal, semi_formal, simple"""

INTERVIEW_PROMPT = """You are an experienced career coach helping engineering students prepare for job interviews.

Question Type: {question_type}
Candidate Background: {background}

Task: Generate a strong, natural-sounding interview answer that:
1. Sounds authentic and personal (not generic or templated)
2. Highlights relevant skills and experiences
3. Follows STAR method where applicable (Situation, Task, Action, Result)
4. Is appropriate length (1-2 minutes when spoken)
5. Avoids clichés and buzzwords

Provide a well-structured answer with key points and suggested delivery notes."""

def call_gemini(prompt: str, emergent_mode: bool = False) -> str:
    """Call Gemini API with appropriate settings"""
    # Use same model but different config for emergent mode (faster, more focused)
    model_name = "gemini-2.5-flash"
    temperature = 0.3 if emergent_mode else 0.7
    max_tokens = 1500 if emergent_mode else 2000
    
    response = client.models.generate_content(
        model=model_name,
        contents=prompt,
        config=genai.types.GenerateContentConfig(
            temperature=temperature,
            max_output_tokens=max_tokens
        )
    )
    return response.text

def generate_pdf(content: Dict[str, Any], doc_type: str) -> BytesIO:
    """Generate PDF document"""
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=24, textColor=colors.HexColor('#7C3AED'), spaceAfter=30)
    heading_style = ParagraphStyle('CustomHeading', parent=styles['Heading2'], fontSize=16, textColor=colors.HexColor('#4F46E5'), spaceAfter=12)
    
    story = []
    
    if doc_type == "resume":
        story.append(Paragraph("Professional Resume", title_style))
        story.append(Paragraph(content.get('summary', ''), styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        
        if content.get('skills'):
            story.append(Paragraph("Skills", heading_style))
            story.append(Paragraph(", ".join(content['skills']), styles['Normal']))
            story.append(Spacer(1, 0.2*inch))
    
    elif doc_type == "project":
        story.append(Paragraph("Final Year Project Report", title_style))
        story.append(Paragraph("Abstract", heading_style))
        story.append(Paragraph(content.get('abstract', ''), styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
    
    doc.build(story)
    buffer.seek(0)
    return buffer

def generate_docx(content: Dict[str, Any], doc_type: str) -> BytesIO:
    """Generate DOCX document"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    if doc_type == "resume":
        title = doc.add_heading('Professional Resume', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(content.get('summary', ''))
        doc.add_heading('Skills', level=1)
        doc.add_paragraph(', '.join(content.get('skills', [])))
        
        if content.get('experience'):
            doc.add_heading('Experience', level=1)
            for exp in content['experience']:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"{exp.get('title', '')} - {exp.get('company', '')} ({exp.get('duration', '')})").bold = True
    
    elif doc_type == "project":
        title = doc.add_heading('Final Year Project Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_heading('Abstract', level=1)
        doc.add_paragraph(content.get('abstract', ''))
        
        doc.add_heading('Problem Statement', level=1)
        doc.add_paragraph(content.get('problem_statement', ''))
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Auth endpoints
@api_router.post("/auth/signup", response_model=AuthResponse)
async def signup(user_data: UserSignup):
    existing = await db.users.find_one({"email": user_data.email}, {"_id": 0})
    if existing:
        raise HTTPException(status_code=400, detail="Email already registered")
    
    user = User(
        name=user_data.name,
        email=user_data.email,
        credits=100
    )
    
    user_dict = user.model_dump()
    user_dict['password'] = get_password_hash(user_data.password)
    user_dict['created_at'] = user_dict['created_at'].isoformat()
    
    await db.users.insert_one(user_dict)
    
    token = create_access_token({"sub": user.id})
    return AuthResponse(token=token, user=user)

@api_router.post("/auth/login", response_model=AuthResponse)
async def login(credentials: UserLogin):
    user_doc = await db.users.find_one({"email": credentials.email}, {"_id": 0})
    if not user_doc:
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    if not verify_password(credentials.password, user_doc['password']):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    if isinstance(user_doc.get('created_at'), str):
        user_doc['created_at'] = datetime.fromisoformat(user_doc['created_at'])
    
    user = User(**{k: v for k, v in user_doc.items() if k != 'password'})
    token = create_access_token({"sub": user.id})
    return AuthResponse(token=token, user=user)

@api_router.get("/auth/me", response_model=User)
async def get_me(request: Request, current_user: User = Depends(get_current_user)):
    return current_user

@api_router.post("/auth/session")
async def create_session_from_oauth(request: Request):
    """Exchange Emergent OAuth session_id for user data and create session"""
    body = await request.json()
    session_id = body.get("session_id")
    
    if not session_id:
        raise HTTPException(status_code=400, detail="session_id required")
    
    # Call Emergent Auth API
    import httpx
    async with httpx.AsyncClient() as client:
        response = await client.get(
            "https://demobackend.emergentagent.com/auth/v1/env/oauth/session-data",
            headers={"X-Session-ID": session_id}
        )
        
        if response.status_code != 200:
            raise HTTPException(status_code=400, detail="Invalid session_id")
        
        oauth_data = response.json()
    
    # Check if user exists by email
    user_doc = await db.users.find_one({"email": oauth_data["email"]}, {"_id": 0})
    
    if user_doc:
        # Update existing user
        await db.users.update_one(
            {"email": oauth_data["email"]},
            {"$set": {
                "name": oauth_data["name"],
                "picture": oauth_data.get("picture")
            }}
        )
        user_id = user_doc["id"]
    else:
        # Create new user with 100 free credits
        user = User(
            name=oauth_data["name"],
            email=oauth_data["email"],
            credits=100
        )
        user_dict = user.model_dump()
        user_dict['picture'] = oauth_data.get("picture")
        user_dict['created_at'] = user_dict['created_at'].isoformat()
        await db.users.insert_one(user_dict)
        user_id = user.id
    
    # Create session in database
    session_token = oauth_data["session_token"]
    expires_at = datetime.now(timezone.utc) + timedelta(days=7)
    
    await db.user_sessions.insert_one({
        "user_id": user_id,
        "session_token": session_token,
        "expires_at": expires_at.isoformat(),
        "created_at": datetime.now(timezone.utc).isoformat()
    })
    
    # Get updated user
    user_doc = await db.users.find_one({"id": user_id}, {"_id": 0})
    if isinstance(user_doc.get('created_at'), str):
        user_doc['created_at'] = datetime.fromisoformat(user_doc['created_at'])
    
    user_response = User(**user_doc)
    
    # Set httpOnly cookie
    response = {"user": user_response, "session_token": session_token}
    
    from fastapi.responses import JSONResponse
    json_response = JSONResponse(content={"user": user_response.model_dump(), "session_token": session_token})
    json_response.set_cookie(
        key="session_token",
        value=session_token,
        httponly=True,
        secure=True,
        samesite="none",
        path="/",
        max_age=7*24*60*60  # 7 days
    )
    
    return json_response

@api_router.post("/auth/logout")
async def logout(request: Request, current_user: User = Depends(get_current_user)):
    """Logout user and clear session"""
    session_token = request.cookies.get("session_token")
    
    if session_token:
        await db.user_sessions.delete_one({"session_token": session_token})
    
    from fastapi.responses import JSONResponse
    json_response = JSONResponse(content={"success": True})
    json_response.delete_cookie(key="session_token", path="/")
    
    return json_response

# AI Generation endpoints
@api_router.post("/generate/resume")
async def generate_resume(request: ResumeRequest, current_user: User = Depends(get_current_user)):
    credits_needed = int(CREDIT_COSTS["resume"] * (1.3 if request.emergent_mode else 1))
    
    if current_user.credits < credits_needed:
        raise HTTPException(status_code=402, detail="Insufficient credits")
    
    prompt = RESUME_PROMPT.format(
        target_role=request.target_role,
        country=request.country,
        resume_text=request.resume_text
    )
    
    try:
        result = call_gemini(prompt, request.emergent_mode)
        
        import json
        try:
            result_json = json.loads(result.strip().replace('```json', '').replace('```', ''))
        except:
            result_json = {"summary": result, "skills": [], "experience": [], "education": [], "projects": []}
        
        await db.users.update_one(
            {"id": current_user.id},
            {"$inc": {"credits": -credits_needed}}
        )
        
        history = GenerationHistory(
            user_id=current_user.id,
            type="Resume",
            title=f"{request.target_role} - {request.country}",
            content=result_json,
            credits_used=credits_needed
        )
        history_dict = history.model_dump()
        history_dict['created_at'] = history_dict['created_at'].isoformat()
        await db.generation_history.insert_one(history_dict)
        
        return {"success": True, "data": result_json, "credits_used": credits_needed, "history_id": history.id}
    
    except Exception as e:
        logger.error(f"Resume generation error: {e}")
        raise HTTPException(status_code=500, detail="Generation failed")

@api_router.post("/generate/project")
async def generate_project(request: ProjectRequest, current_user: User = Depends(get_current_user)):
    credits_needed = int(CREDIT_COSTS["project"] * (1.3 if request.emergent_mode else 1))
    
    if current_user.credits < credits_needed:
        raise HTTPException(status_code=402, detail="Insufficient credits")
    
    prompt = PROJECT_PROMPT.format(topic=request.topic, branch=request.branch)
    
    try:
        result = call_gemini(prompt, request.emergent_mode)
        
        import json
        try:
            result_json = json.loads(result.strip().replace('```json', '').replace('```', ''))
        except:
            result_json = {"abstract": result, "problem_statement": "", "objectives": [], "modules": []}
        
        await db.users.update_one(
            {"id": current_user.id},
            {"$inc": {"credits": -credits_needed}}
        )
        
        history = GenerationHistory(
            user_id=current_user.id,
            type="Project",
            title=request.topic,
            content=result_json,
            credits_used=credits_needed
        )
        history_dict = history.model_dump()
        history_dict['created_at'] = history_dict['created_at'].isoformat()
        await db.generation_history.insert_one(history_dict)
        
        return {"success": True, "data": result_json, "credits_used": credits_needed, "history_id": history.id}
    
    except Exception as e:
        logger.error(f"Project generation error: {e}")
        raise HTTPException(status_code=500, detail="Generation failed")

@api_router.post("/generate/english")
async def generate_english(request: EnglishRequest, current_user: User = Depends(get_current_user)):
    credits_needed = int(CREDIT_COSTS["english"] * (1.3 if request.emergent_mode else 1))
    
    if current_user.credits < credits_needed:
        raise HTTPException(status_code=402, detail="Insufficient credits")
    
    prompt = ENGLISH_PROMPT.format(text=request.text)
    
    try:
        result = call_gemini(prompt, request.emergent_mode)
        
        import json
        try:
            result_json = json.loads(result.strip().replace('```json', '').replace('```', ''))
        except:
            result_json = {"formal": result, "semi_formal": result, "simple": result}
        
        await db.users.update_one(
            {"id": current_user.id},
            {"$inc": {"credits": -credits_needed}}
        )
        
        history = GenerationHistory(
            user_id=current_user.id,
            type="English",
            title="Text Improvement",
            content=result_json,
            credits_used=credits_needed
        )
        history_dict = history.model_dump()
        history_dict['created_at'] = history_dict['created_at'].isoformat()
        await db.generation_history.insert_one(history_dict)
        
        return {"success": True, "data": result_json, "credits_used": credits_needed, "history_id": history.id}
    
    except Exception as e:
        logger.error(f"English generation error: {e}")
        raise HTTPException(status_code=500, detail="Generation failed")

@api_router.post("/generate/interview")
async def generate_interview(request: InterviewRequest, current_user: User = Depends(get_current_user)):
    credits_needed = int(CREDIT_COSTS["interview"] * (1.3 if request.emergent_mode else 1))
    
    if current_user.credits < credits_needed:
        raise HTTPException(status_code=402, detail="Insufficient credits")
    
    prompt = INTERVIEW_PROMPT.format(
        question_type=request.question_type,
        background=request.background or "Fresh graduate in computer science"
    )
    
    try:
        result = call_gemini(prompt, request.emergent_mode)
        
        await db.users.update_one(
            {"id": current_user.id},
            {"$inc": {"credits": -credits_needed}}
        )
        
        history = GenerationHistory(
            user_id=current_user.id,
            type="Interview",
            title=request.question_type,
            content={"answer": result},
            credits_used=credits_needed
        )
        history_dict = history.model_dump()
        history_dict['created_at'] = history_dict['created_at'].isoformat()
        await db.generation_history.insert_one(history_dict)
        
        return {"success": True, "data": {"answer": result}, "credits_used": credits_needed, "history_id": history.id}
    
    except Exception as e:
        logger.error(f"Interview generation error: {e}")
        raise HTTPException(status_code=500, detail="Generation failed")

# Document download endpoints
@api_router.get("/download/{history_id}/{format}")
async def download_document(history_id: str, format: str, current_user: User = Depends(get_current_user)):
    history_doc = await db.generation_history.find_one({"id": history_id, "user_id": current_user.id}, {"_id": 0})
    
    if not history_doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    doc_type = history_doc['type'].lower()
    content = history_doc['content']
    
    if format == "pdf":
        buffer = generate_pdf(content, doc_type)
        return StreamingResponse(buffer, media_type="application/pdf", headers={"Content-Disposition": f"attachment; filename={doc_type}_{history_id}.pdf"})
    elif format == "docx":
        buffer = generate_docx(content, doc_type)
        return StreamingResponse(buffer, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f"attachment; filename={doc_type}_{history_id}.docx"})
    else:
        raise HTTPException(status_code=400, detail="Invalid format")

# History endpoints
@api_router.get("/history")
async def get_history(current_user: User = Depends(get_current_user)):
    history_docs = await db.generation_history.find({"user_id": current_user.id}, {"_id": 0}).sort("created_at", -1).limit(50).to_list(50)
    
    for doc in history_docs:
        if isinstance(doc.get('created_at'), str):
            doc['created_at'] = datetime.fromisoformat(doc['created_at'])
    
    return history_docs

# Payment endpoints - Razorpay
@api_router.post("/payments/razorpay/create-order", response_model=RazorpayOrderResponse)
async def create_razorpay_order(request: PurchaseRequest, current_user: User = Depends(get_current_user)):
    package = CREDIT_PACKAGES.get(request.package_id)
    if not package or package['currency'] != 'INR':
        raise HTTPException(status_code=400, detail="Invalid package")
    
    amount_paise = int(package['price'] * 100)
    
    try:
        order = razorpay_client.order.create({
            "amount": amount_paise,
            "currency": "INR",
            "payment_capture": 1
        })
        
        transaction = PaymentTransaction(
            user_id=current_user.id,
            session_id=order['id'],
            amount=package['price'],
            currency='INR',
            package_id=request.package_id,
            credits=package['credits'],
            status='pending',
            payment_status='initiated',
            provider='razorpay'
        )
        transaction_dict = transaction.model_dump()
        transaction_dict['created_at'] = transaction_dict['created_at'].isoformat()
        await db.payment_transactions.insert_one(transaction_dict)
        
        return RazorpayOrderResponse(
            order_id=order['id'],
            amount=amount_paise,
            currency='INR',
            key_id=os.environ['RAZORPAY_KEY_ID']
        )
    except Exception as e:
        logger.error(f"Razorpay order creation failed: {e}")
        raise HTTPException(status_code=500, detail="Order creation failed")

@api_router.post("/payments/razorpay/verify")
async def verify_razorpay_payment(request: RazorpayVerifyRequest, current_user: User = Depends(get_current_user)):
    try:
        razorpay_client.utility.verify_payment_signature({
            'razorpay_order_id': request.order_id,
            'razorpay_payment_id': request.payment_id,
            'razorpay_signature': request.signature
        })
        
        transaction_doc = await db.payment_transactions.find_one(
            {"session_id": request.order_id, "user_id": current_user.id, "payment_status": {"$ne": "paid"}},
            {"_id": 0}
        )
        
        if not transaction_doc:
            raise HTTPException(status_code=404, detail="Transaction not found or already processed")
        
        await db.payment_transactions.update_one(
            {"id": transaction_doc['id']},
            {"$set": {"payment_id": request.payment_id, "status": "completed", "payment_status": "paid"}}
        )
        
        await db.users.update_one(
            {"id": current_user.id},
            {"$inc": {"credits": transaction_doc['credits']}}
        )
        
        return {"success": True, "credits_added": transaction_doc['credits']}
    
    except Exception as e:
        logger.error(f"Razorpay verification failed: {e}")
        raise HTTPException(status_code=400, detail="Payment verification failed")

# Payment endpoints - Stripe
@api_router.post("/payments/stripe/create-checkout", response_model=CheckoutSessionResponse)
async def create_stripe_checkout(request: StripeCheckoutRequest, http_request: Request, current_user: User = Depends(get_current_user)):
    package = CREDIT_PACKAGES.get(request.package_id)
    if not package or package['currency'] != 'USD':
        raise HTTPException(status_code=400, detail="Invalid package")
    
    try:
        host_url = str(http_request.base_url).rstrip('/')
        webhook_url = f"{host_url}/api/webhook/stripe"
        
        stripe_checkout = StripeCheckout(api_key=os.environ['STRIPE_API_KEY'], webhook_url=webhook_url)
        
        success_url = f"{host_url}/purchase/success?session_id={{CHECKOUT_SESSION_ID}}"
        cancel_url = f"{host_url}/purchase"
        
        checkout_request = CheckoutSessionRequest(
            amount=package['price'],
            currency="usd",
            success_url=success_url,
            cancel_url=cancel_url,
            metadata={
                "user_id": current_user.id,
                "package_id": request.package_id,
                "credits": str(package['credits'])
            }
        )
        
        session = await stripe_checkout.create_checkout_session(checkout_request)
        
        transaction = PaymentTransaction(
            user_id=current_user.id,
            session_id=session.session_id,
            amount=package['price'],
            currency='USD',
            package_id=request.package_id,
            credits=package['credits'],
            status='pending',
            payment_status='initiated',
            provider='stripe'
        )
        transaction_dict = transaction.model_dump()
        transaction_dict['created_at'] = transaction_dict['created_at'].isoformat()
        await db.payment_transactions.insert_one(transaction_dict)
        
        return session
    
    except Exception as e:
        logger.error(f"Stripe checkout creation failed: {e}")
        raise HTTPException(status_code=500, detail="Checkout creation failed")

@api_router.get("/payments/stripe/status/{session_id}", response_model=CheckoutStatusResponse)
async def get_stripe_status(session_id: str, current_user: User = Depends(get_current_user)):
    try:
        webhook_url = "dummy"  # Not needed for status check
        stripe_checkout = StripeCheckout(api_key=os.environ['STRIPE_API_KEY'], webhook_url=webhook_url)
        
        checkout_status = await stripe_checkout.get_checkout_status(session_id)
        
        if checkout_status.payment_status == "paid":
            transaction_doc = await db.payment_transactions.find_one(
                {"session_id": session_id, "user_id": current_user.id, "payment_status": {"$ne": "paid"}},
                {"_id": 0}
            )
            
            if transaction_doc:
                await db.payment_transactions.update_one(
                    {"id": transaction_doc['id']},
                    {"$set": {"status": "completed", "payment_status": "paid"}}
                )
                
                await db.users.update_one(
                    {"id": current_user.id},
                    {"$inc": {"credits": transaction_doc['credits']}}
                )
        
        return checkout_status
    
    except Exception as e:
        logger.error(f"Stripe status check failed: {e}")
        raise HTTPException(status_code=500, detail="Status check failed")

@api_router.post("/webhook/stripe")
async def stripe_webhook(request: Request):
    try:
        body = await request.body()
        signature = request.headers.get("Stripe-Signature")
        
        webhook_url = "dummy"
        stripe_checkout = StripeCheckout(api_key=os.environ['STRIPE_API_KEY'], webhook_url=webhook_url)
        
        webhook_response = await stripe_checkout.handle_webhook(body, signature)
        
        logger.info(f"Stripe webhook received: {webhook_response.event_type}")
        
        return {"status": "received"}
    
    except Exception as e:
        logger.error(f"Stripe webhook error: {e}")
        return {"status": "error"}

# Include router
app.include_router(api_router)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
