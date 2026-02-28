from io import BytesIO
from typing import Dict, Any
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_pdf(content: Dict[str, Any], doc_type: str) -> BytesIO:
    """Generate a PDF document from generation content."""
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=letter,
        rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=36
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle', parent=styles['Heading1'],
        fontSize=22, textColor=colors.HexColor('#7C3AED'), spaceAfter=24
    )
    heading_style = ParagraphStyle(
        'CustomHeading', parent=styles['Heading2'],
        fontSize=14, textColor=colors.HexColor('#4F46E5'), spaceAfter=10, spaceBefore=16
    )
    body_style = styles['Normal']

    story = []

    if doc_type == "resume":
        story.append(Paragraph("Professional Resume", title_style))
        story.append(Spacer(1, 0.1 * inch))

        if content.get('summary'):
            story.append(Paragraph("Professional Summary", heading_style))
            story.append(Paragraph(str(content['summary']), body_style))
            story.append(Spacer(1, 0.15 * inch))

        if content.get('skills'):
            story.append(Paragraph("Skills", heading_style))
            story.append(Paragraph(", ".join(content['skills']), body_style))
            story.append(Spacer(1, 0.15 * inch))

        if content.get('experience'):
            story.append(Paragraph("Experience", heading_style))
            for exp in content['experience']:
                title_text = f"<b>{exp.get('title', '')} - {exp.get('company', '')}</b> ({exp.get('duration', '')})"
                story.append(Paragraph(title_text, body_style))
                for point in exp.get('points', []):
                    story.append(Paragraph(f"  - {point}", body_style))
                story.append(Spacer(1, 0.1 * inch))

        if content.get('education'):
            story.append(Paragraph("Education", heading_style))
            for edu in content['education']:
                story.append(Paragraph(
                    f"<b>{edu.get('degree', '')}</b> - {edu.get('institution', '')} ({edu.get('year', '')})",
                    body_style
                ))
            story.append(Spacer(1, 0.15 * inch))

        if content.get('projects'):
            story.append(Paragraph("Projects", heading_style))
            for proj in content['projects']:
                story.append(Paragraph(f"<b>{proj.get('name', '')}</b>", body_style))
                story.append(Paragraph(str(proj.get('description', '')), body_style))
                if proj.get('tech'):
                    story.append(Paragraph(f"Tech: {', '.join(proj['tech'])}", body_style))
                story.append(Spacer(1, 0.1 * inch))

    elif doc_type == "project":
        story.append(Paragraph("Final Year Project Report", title_style))
        story.append(Spacer(1, 0.1 * inch))

        sections = [
            ("Abstract", "abstract"),
            ("Problem Statement", "problem_statement"),
            ("System Architecture", "architecture"),
            ("Implementation", "implementation"),
        ]
        for heading, key in sections:
            if content.get(key):
                story.append(Paragraph(heading, heading_style))
                story.append(Paragraph(str(content[key]), body_style))
                story.append(Spacer(1, 0.15 * inch))

        if content.get('objectives'):
            story.append(Paragraph("Objectives", heading_style))
            for obj in content['objectives']:
                story.append(Paragraph(f"- {obj}", body_style))
            story.append(Spacer(1, 0.15 * inch))

        if content.get('modules'):
            story.append(Paragraph("Modules", heading_style))
            for mod in content['modules']:
                story.append(Paragraph(f"<b>{mod.get('name', '')}</b>: {mod.get('description', '')}", body_style))
            story.append(Spacer(1, 0.15 * inch))

        if content.get('tech_stack'):
            story.append(Paragraph("Technology Stack", heading_style))
            for tech in content['tech_stack']:
                if isinstance(tech, dict):
                    story.append(Paragraph(f"- <b>{tech.get('name', '')}</b>: {tech.get('purpose', '')}", body_style))
                else:
                    story.append(Paragraph(f"- {tech}", body_style))
            story.append(Spacer(1, 0.15 * inch))

        if content.get('future_scope'):
            story.append(Paragraph("Future Scope", heading_style))
            for item in content['future_scope']:
                story.append(Paragraph(f"- {item}", body_style))
            story.append(Spacer(1, 0.15 * inch))

        if content.get('viva_questions'):
            story.append(Paragraph("Viva Questions", heading_style))
            for i, qa in enumerate(content['viva_questions'], 1):
                if isinstance(qa, dict):
                    story.append(Paragraph(f"<b>Q{i}: {qa.get('question', '')}</b>", body_style))
                    story.append(Paragraph(f"A: {qa.get('answer', '')}", body_style))
                    story.append(Spacer(1, 0.1 * inch))

    elif doc_type == "english":
        story.append(Paragraph("English Improvement", title_style))
        for tone, label in [("formal", "Formal"), ("semi_formal", "Semi-Formal"), ("simple", "Simple")]:
            if content.get(tone):
                story.append(Paragraph(label, heading_style))
                story.append(Paragraph(str(content[tone]), body_style))
                story.append(Spacer(1, 0.15 * inch))

    elif doc_type == "interview":
        story.append(Paragraph("Interview Preparation", title_style))
        if content.get('answer'):
            story.append(Paragraph(str(content['answer']), body_style))

    doc.build(story)
    buffer.seek(0)
    return buffer


def generate_docx(content: Dict[str, Any], doc_type: str) -> BytesIO:
    """Generate a DOCX document from generation content."""
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    if doc_type == "resume":
        title = doc.add_heading('Professional Resume', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if content.get('summary'):
            doc.add_heading('Professional Summary', level=1)
            doc.add_paragraph(str(content['summary']))

        if content.get('skills'):
            doc.add_heading('Skills', level=1)
            doc.add_paragraph(', '.join(content['skills']))

        if content.get('experience'):
            doc.add_heading('Experience', level=1)
            for exp in content['experience']:
                p = doc.add_paragraph()
                run = p.add_run(f"{exp.get('title', '')} - {exp.get('company', '')} ({exp.get('duration', '')})")
                run.bold = True
                for point in exp.get('points', []):
                    doc.add_paragraph(point, style='List Bullet')

        if content.get('education'):
            doc.add_heading('Education', level=1)
            for edu in content['education']:
                p = doc.add_paragraph()
                run = p.add_run(f"{edu.get('degree', '')} - {edu.get('institution', '')} ({edu.get('year', '')})")
                run.bold = True

        if content.get('projects'):
            doc.add_heading('Projects', level=1)
            for proj in content['projects']:
                p = doc.add_paragraph()
                run = p.add_run(proj.get('name', ''))
                run.bold = True
                doc.add_paragraph(str(proj.get('description', '')))
                if proj.get('tech'):
                    doc.add_paragraph(f"Technologies: {', '.join(proj['tech'])}")

    elif doc_type == "project":
        title = doc.add_heading('Final Year Project Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        sections = [
            ("Abstract", "abstract"),
            ("Problem Statement", "problem_statement"),
            ("System Architecture", "architecture"),
            ("Implementation", "implementation"),
        ]
        for heading, key in sections:
            if content.get(key):
                doc.add_heading(heading, level=1)
                doc.add_paragraph(str(content[key]))

        if content.get('objectives'):
            doc.add_heading('Objectives', level=1)
            for obj in content['objectives']:
                doc.add_paragraph(obj, style='List Bullet')

        if content.get('modules'):
            doc.add_heading('Modules', level=1)
            for mod in content['modules']:
                p = doc.add_paragraph()
                run = p.add_run(mod.get('name', ''))
                run.bold = True
                p.add_run(f": {mod.get('description', '')}")

        if content.get('tech_stack'):
            doc.add_heading('Technology Stack', level=1)
            for tech in content['tech_stack']:
                if isinstance(tech, dict):
                    doc.add_paragraph(f"{tech.get('name', '')}: {tech.get('purpose', '')}", style='List Bullet')
                else:
                    doc.add_paragraph(str(tech), style='List Bullet')

        if content.get('future_scope'):
            doc.add_heading('Future Scope', level=1)
            for item in content['future_scope']:
                doc.add_paragraph(item, style='List Bullet')

        if content.get('viva_questions'):
            doc.add_heading('Viva Questions', level=1)
            for i, qa in enumerate(content['viva_questions'], 1):
                if isinstance(qa, dict):
                    p = doc.add_paragraph()
                    run = p.add_run(f"Q{i}: {qa.get('question', '')}")
                    run.bold = True
                    doc.add_paragraph(f"A: {qa.get('answer', '')}")

    elif doc_type == "english":
        title = doc.add_heading('English Improvement', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for tone, label in [("formal", "Formal"), ("semi_formal", "Semi-Formal"), ("simple", "Simple")]:
            if content.get(tone):
                doc.add_heading(label, level=1)
                doc.add_paragraph(str(content[tone]))

    elif doc_type == "interview":
        title = doc.add_heading('Interview Preparation', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if content.get('answer'):
            doc.add_paragraph(str(content['answer']))

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
